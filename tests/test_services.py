# tests/test_services.py

import pytest
import io
import openpyxl
import requests
from unittest.mock import patch, MagicMock
from docx import Document

from app.services import document_service
from app.services import graph_service
from app.services.graph_service import GraphAPIError


class TestDocumentService:
    """Тесты для сервиса генерации Word-документов."""

    def test_generate_word_from_data(self):
        """
        Тест: Проверяет, что плейсхолдеры в шаблоне Word корректно заменяются.
        """
        # 1. Создаем "моковый" шаблон Word в памяти
        doc = Document()
        doc.add_paragraph("Здравствуйте, {{ИМЯ}}!")
        doc.add_paragraph("Добро пожаловать в город {{ГОРОД}}.")
        doc.add_paragraph("Этот текст останется без изменений.")
        
        table = doc.add_table(rows=1, cols=2)
        cell1 = table.cell(0, 0)
        cell1.text = "Ключ: {{КЛЮЧ}}"
        cell2 = table.cell(0, 1)
        cell2.text = "Еще один ключ: {{КЛЮЧ}}"
        
        template_stream = io.BytesIO()
        doc.save(template_stream)
        template_stream.seek(0)

        # 2. Определяем данные для замены
        placeholders = {
            "{{ИМЯ}}": "Иван",
            "{{ГОРОД}}": "Москва",
            "{{КЛЮЧ}}": "ЗНАЧЕНИЕ"
        }

        # 3. Вызываем тестируемую функцию
        result_stream = document_service.generate_word_from_data(template_stream, placeholders)

        # 4. Проверяем результат
        result_doc = Document(result_stream)
        
        assert "Здравствуйте, Иван!" in [p.text for p in result_doc.paragraphs]
        assert "Добро пожаловать в город Москва." in [p.text for p in result_doc.paragraphs]
        assert "Этот текст останется без изменений." in [p.text for p in result_doc.paragraphs]
        
        result_table = result_doc.tables[0]
        assert result_table.cell(0, 0).text == "Ключ: ЗНАЧЕНИЕ"
        assert result_table.cell(0, 1).text == "Еще один ключ: ЗНАЧЕНИЕ"


class TestGraphService:
    """Тесты для сервиса работы с Excel-файлами (парсинг)."""

    def test_read_row_from_excel_bytes_success(self):
        """
        Тест: Проверяет корректное чтение строки из Excel-файла
        и преобразование ее в словарь плейсхолдеров.
        """
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet["A1"] = " № "
        sheet["B1"] = "Наименование изделия"
        sheet["C1"] = "Количество"
        sheet["A2"] = 1
        sheet["B2"] = "Крышка"
        sheet["C2"] = 15.5
        
        excel_stream = io.BytesIO()
        workbook.save(excel_stream)
        excel_bytes = excel_stream.getvalue()

        placeholders = graph_service.read_row_from_excel_bytes(excel_bytes, row_number=2)

        expected_placeholders = {
            "{{№}}": "1",
            "{{Наименование изделия}}": "Крышка",
            "{{Количество}}": "15.5"
        }
        assert placeholders == expected_placeholders

    def test_read_row_from_excel_bytes_invalid_row(self):
        """
        Тест: Проверяет, что функция вызывает исключение при неверном номере строки.
        """
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet["A1"] = "Header"
        sheet["A2"] = "Data"
        
        excel_stream = io.BytesIO()
        workbook.save(excel_stream)
        excel_bytes = excel_stream.getvalue()

        with pytest.raises(IndexError):
            graph_service.read_row_from_excel_bytes(excel_bytes, row_number=3)
        
        with pytest.raises(IndexError):
            graph_service.read_row_from_excel_bytes(excel_bytes, row_number=1)


class TestGraphServiceWithMocks:
    """
    НОВЫЙ КЛАСС. Тесты для сетевой части graph_service с использованием "моков".
    """

    @patch('app.services.graph_service.requests.post')
    def test_get_access_token_success(self, mock_post, monkeypatch):
        """Тест: Проверяет успешное получение токена доступа."""
        monkeypatch.setenv("MS_CLIENT_ID", "test_client_id")
        monkeypatch.setenv("MS_CLIENT_SECRET", "test_client_secret")
        monkeypatch.setenv("MS_TENANT_ID", "test_tenant_id")
        
        mock_response = MagicMock()
        mock_response.json.return_value = {'access_token': 'fake_token'}
        mock_response.raise_for_status.return_value = None
        mock_post.return_value = mock_response

        token = graph_service._get_access_token()
        assert token == 'fake_token'

    @patch('app.services.graph_service.requests.post')
    def test_get_access_token_failure(self, mock_post, monkeypatch):
        """Тест: Проверяет обработку ошибки при получении токена."""
        monkeypatch.setenv("MS_CLIENT_ID", "test_client_id")
        monkeypatch.setenv("MS_CLIENT_SECRET", "test_client_secret")
        monkeypatch.setenv("MS_TENANT_ID", "test_tenant_id")

        mock_response = MagicMock()
        mock_response.raise_for_status.side_effect = requests.exceptions.HTTPError("401 Unauthorized")
        mock_post.return_value = mock_response
        
        with pytest.raises(GraphAPIError) as excinfo:
            graph_service._get_access_token()
        assert "Ошибка сети при получении токена доступа" in str(excinfo.value)

    @patch('app.services.graph_service._get_access_token')
    @patch('app.services.graph_service.requests.get')
    def test_download_file_from_onedrive_success(self, mock_get, mock_get_token, monkeypatch):
        """Тест: Проверяет успешное скачивание файла."""
        monkeypatch.setenv("MS_ONEDRIVE_USER_ID", "test_user_id")
        mock_get_token.return_value = 'fake_token'
        
        mock_response = MagicMock()
        mock_response.content = b'excel file content'
        mock_response.raise_for_status.return_value = None
        mock_get.return_value = mock_response
        
        file_content = graph_service.download_file_from_onedrive('/test.xlsx')
        assert file_content == b'excel file content'

    @patch('app.services.graph_service._get_access_token')
    @patch('app.services.graph_service.requests.get')
    def test_download_file_from_onedrive_not_found(self, mock_get, mock_get_token, monkeypatch):
        """Тест: Проверяет обработку ошибки 404 (файл не найден)."""
        monkeypatch.setenv("MS_ONEDRIVE_USER_ID", "test_user_id")
        mock_get_token.return_value = 'fake_token'
        
        mock_response = MagicMock()
        mock_response.status_code = 404
        mock_get.return_value = mock_response
        
        with pytest.raises(FileNotFoundError) as excinfo:
            graph_service.download_file_from_onedrive('/not_found.xlsx')
        assert "Файл не найден в OneDrive" in str(excinfo.value)